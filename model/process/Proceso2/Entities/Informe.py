import os
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement, ns

class Informe():
    def __init__(self, path=None):
        if not path:
            self.document = Document()
            currentdir = os.path.dirname(os.path.realpath(__file__))
            self.document.add_picture(currentdir +"/imagen.png",width=Inches(5.9),height=Inches(1))
            self.add_page_number(self.document.sections[0].footer.paragraphs[0].add_run("\t"))
            self.document.styles["Footer"]
        else:
            self.document = Document(path)

    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(self, run):
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def devolver_documento(self):
        if self.document:
            return self.document
        return None

    def insertar_titulo(self, titulo):
        try:
            self.document.add_heading(titulo, 0)
        except Exception as e:
            print("Problema al insertar la cabecera: " + str(e))

    def guardar_documento(self, path):
        self.document.save(path)

    def salto_de_pagina(self):
        self.document.add_page_break()

    # def insertar_parrafo(self, text, contenido=None):
    #     return self.document.add_paragraph(text).add_run(contenido)
    
    # def insertar_linea_parrafo(self, text):
    #     return self.document.add_paragraph().add_run(text)

    
#if __name__ == "__main__":
    #informe = Informe()
    #informe.insertar_titulo("Titulo a poner")
    #informe.guardar_documento("prueba.docx")