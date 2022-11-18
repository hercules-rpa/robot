import time
import os
from datetime import datetime
import json
from model.RPA import RPA
from rpa_robot.ControllerRobot import ControllerRobot
from rpa_robot.ControllerSettings import ControllerSettings
from model.process.UtilsProcess import UtilsProcess
import requests
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from model.process.Process2.Entities.Committee.Committee import Committee
from model.process.Process2.Entities.Committee.CommitteeArquitecturaIngenieriaCivilYUrbanismo import CommitteeArquitecturaIngenieriaCivilYUrbanismo
from model.process.Process2.Entities.Committee.CommitteeBiblioteconomiaYDocumentacion import CommitteeBiblioteconomiaYDocumentacion
from model.process.Process2.Entities.Committee.CommitteeBiologiaCelularYMolecular import CommitteeBiologiaCelularYMolecular
from model.process.Process2.Entities.Committee.CommitteeCienciasBiomedicas import CommitteeCienciasBiomedicas
from model.process.Process2.Entities.Committee.CommitteeCienciasComportamiento import CommitteeCienciasComportamiento
from model.process.Process2.Entities.Committee.CommitteeCienciasEconomicasYEmpresariales import CommitteeCienciasEconomicasYEmpresariales
from model.process.Process2.Entities.Committee.CommitteeCienciasEducacion import CommitteeCienciasEducacion
from model.process.Process2.Entities.Committee.CommitteeCienciasNaturaleza import CommitteeCienciasNaturaleza
from model.process.Process2.Entities.Committee.CommitteeCienciasSociales import CommitteeCienciasSociales
from model.process.Process2.Entities.Committee.CommitteeEstudiosGenero import CommitteeEstudiosGenero
from model.process.Process2.Entities.Committee.CommitteeFilosofiaFilologiaYLinguistica import CommitteeFilosofiaFilologiaYLinguistica
from model.process.Process2.Entities.Committee.CommitteeIngenierias import CommitteeIngenierias
from model.process.Process2.Entities.Committee.CommitteeQuimica import CommitteeQuimica
from model.process.Process2.Entities.Patent import Patent
from model.process.Process2.Entities.Conference import Conference
from model.process.Process2.Entities.Committee.CommitteeAntropologiaSocial import CommitteeAntropologiaSocial
from model.process.Process2.Entities.Accreditation import Accreditation
from model.process.Process2.Entities.Evaluations.SexenioEvaluation import SexenioEvaluation
from model.EDMA import EDMA
from model.process.Process2.Entities.Evaluator import Evaluator
from model.process.Process2.Entities.ResearcherData import ResearcherData
from model.process.Process2.Entities.Enums.ResearcherIdentifyType import ResearcherIdentifyType
from model.process.Process2.Entities.Report import Report
from model.process.Process2.Entities.RO import RO
from model.process.ProcessCommand import ProcessCommand, ProcessID
from model.process.ProcessCommand import Pstatus as pstatus

VERSION = "1.0"
NAME = "Sexenios"
DESCRIPTION = "Proceso que genera un informe de ayuda al investigador para preparar un sexenio"
REQUIREMENTS = ["python-docx", "pandas", "sparqlwrapper", "requests"]
ID = ProcessID.GENERATE_SEXENIO.value
cs = ControllerSettings()

class ProcessSexenios(ProcessCommand):
    def __init__(self, id_schedule, id_log, id_robot, priority, log_file_path, parameters=None, ip_api=None, port_api=None):
        ProcessCommand.__init__(self, ID, NAME, REQUIREMENTS, DESCRIPTION,
                                id_schedule, id_log, id_robot, priority, log_file_path, parameters, ip_api, port_api)
        
        cr = ControllerRobot()        
        self.rpa:RPA = RPA(cr.robot.token)

    def execute(self):
        """
        Método encargado de la ejecución del proceso de sexenios.
        """
        self.state = pstatus.RUNNING
        self.log.start_log(time.time())
        self.log.state = "OK"
        self.notify_update('Comienza la ejecución del proceso.')

        if not self.parameters:
            self.notify_update("No se han establecido parámetros")
            self.log.state = "ERROR"
            self.log.end_log(time.time())
            self.state = pstatus.FINISHED
            return

        self.notify_update('Obteniendo el parámetro "comité".')
        committee: Committee = self.get_committee(self.parameters['comite'])
        if not committee:
            self.notify_update(
                'ERROR no se ha podido obtener el comité.')
            self.log.state = "ERROR"
            self.log.end_log(time.time())
            self.state = pstatus.FINISHED
            return
        else:
            self.notify_update('Obteniendo parámetro "período".')
            period = self.get_period(self.parameters['periodo'])
            self.notify_update('Período: ' + period)

            # comprobar si el comité no usa libros o capitulos de libros
            self.notify_update(
                'Obteniendo parámetros relacionados con el investigador.')
            researcherInfo = self.get_researcher_info(self.parameters)
            
            edma:EDMA = cs.get_edma(self.ip_api, self.port_api)            
            if edma: 
                self.notify_update(
                    'Obteniendo los datos del investigador utilizando ED.')
                researcherData = edma.get_researcher_data(
                    researcherInfo)
                if not researcherData:
                    self.notify_update(
                        'ERROR no se ha podido obtener la información del investigador.')
                    self.log.state = "ERROR"
                    self.log.end_log(time.time())
                    self.state = pstatus.FINISHED
                    return

                self.log.completed = 10

                self.notify_update(
                    "Se procede a recuperar producción científica del investigador de la base de datos de Hércules-EDMA")

                self.notify_update(
                    'Obteniendo la lista de artículos del investigador.')

                articles = self.get_articles(
                    researcherInfo, period, committee.authorship_order, edma)
                chapter_books = None
                books = None
                if committee.books_caps:
                    self.notify_update(
                        'Obteniendo los capítulos de libros del investigador.')
                    chapter_books = edma.get_chapters_books(
                        researcherInfo, period)
                    self.notify_update(
                        'Obteniendo los libros del investigador.')
                    books = edma.get_books(researcherInfo, period)

                conferences: list = None
                if committee.conferences:
                    self.notify_update(
                        'Obteniendo los trabajos presentados en congresos del investigador.')
                    conferences = self.get_conferences(
                        researcherInfo, period, edma)

                patents: list = None
                if committee.patents:
                    self.notify_update(
                        'Obteniendo las patentes del investigador.')
                    patents = self.get_patents(researcherInfo,edma, period)

                self.log.completed = 40

                self.notify_update(
                    "Se procederá a realizar el informe para preparación de sexenios de " + researcherData.name)
                report = Report()
                report.add_title(
                    'Informe de autoevaluación de la investigación para la preparación de un sexenio')
                document = report.get_document()

                self.notify_update(
                    'Insertando la información del investigador en el informe.')

                self.print_researcher_data(
                    document, researcherData, researcherInfo, committee, period)

                self.notify_update(
                    'Obteniendo la baremación en base a los requisitos del comité ' + committee.name)
                evaluation: SexenioEvaluation = committee.get_evaluation_sexenio(
                    articles)
                self.notify_update(
                    'Insertando la baremación obtenida en el informe.')
                self.print_scale(document, evaluation, 1)

                main = []
                substitutes = []
                if articles:
                    if evaluation.main_production:
                        main = evaluation.main_production
                    elif len(articles) > 5:
                        main = articles[0:5]
                    else:
                        main = articles

                    substitutes = evaluation.get_substitute_production()
                    if len(substitutes) > 30:
                        substitutes = substitutes[0:30]

                self.log.completed = 50

                self.notify_update(
                    "Insertando en el informe la producción científica principal.")

                self.print_scientific_production(
                    document, main, committee, 'principal', 2.1)
                self.log.completed = 60

                self.notify_update(
                    "Insertando en el informe la producción científica sustitutoria.")
                self.log.completed = 70
                self.print_scientific_production(
                    document, substitutes, committee, 'sustitutoria', 2.2)

                if committee.authors_alert:
                    self.notify_update(
                        'Insertando en el informe los libros.')
                    self.print_books(document, books, committee, 3)

                    self.notify_update(
                        'Insertando en el informe los capítulos de libros.')
                    self.print_chapter_books(
                        document, chapter_books, committee, 4)

                self.log.completed = 80

                self.print_conferences_patents(document, committee, conferences, patents, 5)

                self.log.completed = 90

                document.add_page_break()

                self.notify_update("Persistiendo informe...")
                doc_name = self.save_report(
                    document, researcherData.name, True)
                self.notify_update(
                    "Informe generado y guardado con el nombre "+doc_name + ".")
                self.log.completed = 90

                self.notify_update("Subiendo informe al CDN...")
                url_cdn = cs.get_url_upload_cdn(self.ip_api, self.port_api)
                response = self.upload_file(doc_name, url_cdn)
                status_send = None
                if response and response.status_code == 200:
                    upload_response = json.loads(response.text)
                    self.notify_update(
                        "Fichero subido correctamente al CDN, url de descarga: "+upload_response["url_cdn"])

                    config = cs.get_globals_settings(self.ip_api, self.port_api)
                    
                    if config:
                        json_dicti = json.loads(config)
                        if json_dicti: 
                            url = json_dicti['edma_host_servicios'] + '/editorcv/Sexenios/Notify'
                            response = self.send_report(upload_response["url_cdn"], researcherInfo[1], url)
                            if response:
                                status_send =response.status_code
                    
                    self.result = upload_response["url_cdn"]

                if status_send and status_send == 200:
                    self.notify_update("Informe notificado correctamente")
                else:
                    self.notify_update("ERROR en la notificación del informe")
                    self.log.state = "ERROR"
                
                
                if os.path.exists(doc_name):
                    os.remove(doc_name)
            else:
                self.notify_update('ERROR en la obtención de parámetros para la consulta a Hércules-EDMA.')
                self.log.state = "ERROR"

        self.notify_update(
            'Proceso de generación de informe de sexenio finalizado.')
        self.log.completed = 100
        self.log.end_log(time.time())
        self.state = pstatus.FINISHED

    def pause(self):
        pass

    def kill(self):
        pass

    def resume(self):
        pass

    def send_report(self, url_cdn, researcher, url):
        if url_cdn and url:
            destination = url + '?url_cdn='+ url_cdn + '&idUsuario=' + researcher
            return requests.post(destination)
        return None


    def get_authors_formatted(self, df) -> str:
        """
        Método que formatea el nombre de los autores.
        :param df propiedad que incluye el nombre de los autores sin formatear
        :return devuelve un string con el nombre de los autores concadenado por comas     
        """
        authors = ''
        for name in df.loc[:, "nombreAutor.value"]:
            authors += name+', '
        return authors[:-2]

    def get_committee(self, id: str) -> Committee:
        """
        Método que crea el comité que evalúa la solicitud del sexenio.
        :param id identificador del comité
        :return Committee objeto que contiene la información del comité
        """
        result: Committee = None
        currentdir = os.path.dirname(os.path.realpath(__file__))
        utils = UtilsProcess()
        conf = utils.get_configurations(
            currentdir+'/Process2/Configurations/comite.json')
        if id == "2":
            result = CommitteeQuimica(id, conf[id])
        elif id == "3":
            result = CommitteeBiologiaCelularYMolecular(id, conf[id])
        elif id == "4":
            result = CommitteeCienciasBiomedicas(id, conf[id])
        elif id == "5":
            result = CommitteeCienciasNaturaleza(id, conf[id])
        elif id == "7":
            result = CommitteeIngenierias(id, conf[id])
        elif id == "8":
            is_tecnology = False
            if self.parameters and 'perfil_tecnologico' in self.parameters:
                is_tecnology = self.parameters['perfil_tecnologico']

            self.notify_update(
                'Perfil tecnológico: ' + str(is_tecnology))
            result = CommitteeArquitecturaIngenieriaCivilYUrbanismo(
                id, conf[id], False, is_tecnology)

        elif id == "9":
            if 'subcomite' in self.parameters:
                sub = str(self.parameters['subcomite'])

                if sub == "1":
                    result = CommitteeCienciasSociales(id, conf[id])
                elif sub == "2" or sub == "6":
                    result = CommitteeCienciasComportamiento(id, conf[id])
                elif sub == "3":
                    result = CommitteeBiblioteconomiaYDocumentacion(id, conf[id])
                elif sub == "4":
                    result = CommitteeEstudiosGenero(id, conf[id])
                elif sub == "5":
                    result = CommitteeAntropologiaSocial(id, conf[id])
        elif id == "10":
            result = CommitteeCienciasEducacion(id, conf[id])
        elif id == "11":
            result = CommitteeCienciasEconomicasYEmpresariales(id, conf[id])
        elif id == "14":
            result = CommitteeFilosofiaFilologiaYLinguistica(id, conf[id])
        return result

    def get_researcher_info(self, parameters):
        """
        Método que obtiene el tipo de identificador del investigador y el identificador
        de los parámetros que recibe el proceso para su ejecución.
        :param parameters conjunto de parámetros específicos del investigador
        :return tuple tupla donde el primer elemento es el identificador del tipo de id y el segundo el identificador      
        """
        typeId: ResearcherIdentifyType = ResearcherIdentifyType.NODEFINIDO
        researcher:str = None
        try:
            if 'investigador' in parameters:
                researcher = parameters['investigador']
                
                if researcher:
                    utils = UtilsProcess()
                    regex = [("^[0-9]{4}-[0-9]{4}-[0-9]{4}-[0-9]{4}$", ResearcherIdentifyType.ORCID),
                    ('^[0-9]{8}$', ResearcherIdentifyType.PERSONAREF),
                    ('^[a-z0-9]+[\._]?[a-z0-9]+[@]\w+[.]\w{2,3}$', ResearcherIdentifyType.EMAIL)]
                    cont = 0
                    while (len(regex)>cont and typeId == ResearcherIdentifyType.NODEFINIDO):
                        if utils.is_match(regex[cont][0], researcher):
                            typeId = regex[cont][1]
                        cont+=1            

                if typeId == ResearcherIdentifyType.NODEFINIDO:
                    researcher = None

        except Exception as e:
            print(e)
            self.notify_update(
                'ERROR al obtener el identificador del investigador.')

        return (typeId, researcher)

    def get_period_array(self, period):
        """
        Método que obtiene calcula el período de años en base a un string.
        Ejemplo: '1,2,5-7' a [1,2,5,6,7]
        :param period periodo a formatear
        :return array que incluye los años que componen el período
        """
        result = []
        try:
            for part in period.split(','):
                if '-' in part:
                    a, b = part.split('-')
                    a, b = int(a), int(b)
                    result.extend(range(a, b + 1))
                else:
                    a = int(part)
                    result.append(a)

        except Exception as e:
            print(e)
            self.notify_update(
                'ERROR al obtener el parámetro "periodo".')
        return result

    def get_period(self, period) -> str:
        """
        Método que obtiene los años que forman parte del período separados por comas.
        :param periodo años que forman parte del período
        :return str cadena de caracteres que separa por años y concatena por comas
        """
        result = self.get_period_array(period)
        period = ''
        if result:
            for a in result:
                period += str(a)+','

            period = period[:-1]
        return period

    def get_article(self, id, edma:EDMA) -> RO:
        """
        Método que obtiene un artículo e hidrata la entidad RO.
        :param id identificador del artículo
        :param edma instancia con la información de Hércules-EDMA
        :return RO objeto que contiene la información del artículo
        """
        article: RO = None
        try:
            article = edma.get_article(id)
            if article:
                authors = edma.get_authors_article(id)
                article.authors = self.get_authors_formatted(authors)
                article.authors_number = len(authors.loc[:, "nombreAutor.value"])
        except:
            self.notify_update(
                'ERROR a obtener el artículo con id: ' + id)

        return article

    def get_patent(self, id, edma:EDMA) -> Patent:
        """
        Método que obtiene una patente e hidrata la entidad Patente.
        :param id identificador de la patente
        :param edma instancia con la información de Hércules-EDMA
        :return Patent objeto que contiene la información de la patente
        """
        patent: Patent = None
        try:
            patent = edma.get_title_date_patent(id)
            if patent:
                authors = edma.get_authors_patent(id)
                patent.authors = self.get_authors_formatted(authors)
        except:
            self.notify_update(
                'ERROR a obtener la patente con id: ' + id)

        return patent

    def get_conference(self, id, edma:EDMA) -> Conference:
        """
        Método que obtiene un trabajo presentado en un congreso e hidrata la entidad Congreso.
        :param id identificador del congreso
        :param edma instancia con la información de Hércules-EDMA
        :return Conference objeto que contiene la información relacionada con el congreso
        """
        conference: Conference = None
        try:
            conference = edma.get_conference(id)
            if conference:
                autores = edma.get_authors_conference_paper(id)
                conference.authors = self.get_authors_formatted(autores)
                conference.num_authors = len(autores.loc[:, "nombreAutor.value"])
        except:
            self.notify_update(
                'ERROR a obtener el trabajo presentado en un congreso con id: ' + id)

        return conference

    def get_patents(self, researcherInfo, edma:EDMA, period:str='') -> list:
        """
        Método que obtiene la lista de patentes de un investigador.
        :param infoInvestigador información del investigador
        :param edma instancia con la información de Hércules-EDMA
        :param periodo período a consultar
        :return lista de patentes
        """
        result = []
        dataframe = edma.get_researcher_patents(researcherInfo, period)
        if not dataframe.empty:
            self.notify_update(
                'Obteniendo la información individual de la lista de patentes.')
            for i in range(len(dataframe)):
                patente = self.get_patent(dataframe.iloc[i]['doc.value'], edma)
                if patente:
                    result.append(patente)
        return result

    def get_articles(self, researcherInfo, period, authorship_order, edma:EDMA, 
    max_article: int = 0) -> list:
        """
        Método que obtiene la lista de artículos de un investigador.
        :param researcherInfo información del investigador
        :param period período a consultar
        :paramm authorship_order define si hay que ordenar el conjunto de elementos por el número de autores
        :param edma instancia con la información de Hércules-EDMA
        :param max_article número máximo de artículos
        :return lista de artículos
        """
        result = []
        dataframe_art = edma.get_researcher_articles(
            researcherInfo, max_article, period, authorship_order)
        if not dataframe_art.empty:
            self.notify_update(
                'Obteniendo la información individual de la lista de artículos.')
            for i in range(len(dataframe_art)):
                article = self.get_article(
                    dataframe_art.iloc[i]['doc.value'], edma)
                if article:
                    article.author_position = dataframe_art.iloc[i]['posicion.value']
                    result.append(article)
        return result

    def get_conferences(self, researcherInfo, period, edma:EDMA) -> list:
        """
        Método que obtiene la lista de trabajos presentados en congresos de un investigador.
        :param researcherInfo información del investigador
        :param period período a consultar
        :param edma instancia con la información de Hércules-EDMA
        :return lista de congresos
        """
        result = []
        dataframe_art = edma.get_conferences(
            researcherInfo, period)
        if not dataframe_art.empty:
            self.notify_update(
                'Obteniendo la información individual de la lista de trabajos presentados en congresos.')
            for i in range(len(dataframe_art)):
                conference = self.get_conference(
                    dataframe_art.iloc[i]['doc.value'], edma)
                if conference:
                    result.append(conference)
        return result

    def print_researcher_data(self, document, researcherData: ResearcherData,
        researcherInfo, evaluator: Evaluator, period: str = '', 
        accreditation: Accreditation = None):
        """
        Método que imprime los datos del investigador.
        :param document documento sobre el que se imprimirá la información
        :param researcherData objeto que contiene los datos del investigador
        :param researcherInfo tupla que contiene información del investigador
        :param evaluator evaluador de la solicitud del informe de sexenio/acreditación
        :param period período del sexenio
        :param accreditation objeto que contiene la información relacionada con la acreditación.
        """
        table = document.add_table(rows=1, cols=2)
        table.style = 'Colorful List Accent 1'
        table.allow_autofit = False
        row = table.rows[0].cells
        d = row[0].add_paragraph('Información del investigador/a')
        row[0].merge(row[1])
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row = table.add_row().cells
        d = row[0].add_paragraph()
        d.add_run('Nombre y apellidos')
        d = row[1].add_paragraph(researcherData.name)

        row = table.add_row().cells
        d = row[0].add_paragraph()
        d.add_run('Universidad')
        d = row[1].add_paragraph(researcherData.university)

        row = table.add_row().cells
        d = row[0].add_paragraph()
        d.add_run('Departamento')
        d = row[1].add_paragraph(researcherData.departament)

        if researcherData.email:
            row = table.add_row().cells
            d = row[0].add_paragraph()
            d.add_run('Email')
            d = row[1].add_paragraph(researcherData.email)

        if researcherInfo[0] == ResearcherIdentifyType.ORCID:
            row = table.add_row().cells
            d = row[0].add_paragraph()
            d.add_run('ORCID')
            d = row[1].add_paragraph(researcherInfo[1])

        if period:
            row = table.add_row().cells
            d = row[0].add_paragraph()
            d.add_run('Período del sexenio')
            d = row[1].add_paragraph(period)

        row = table.add_row().cells
        if evaluator and evaluator.is_commission:
            d = row[0].add_paragraph('Comisión')
        else:
            d = row[0].add_paragraph('Comité')
        d = row[1].add_paragraph(evaluator.name)

        if evaluator.authorship_str:
            row = table.add_row().cells
            d = row[0].add_paragraph()
            d.add_run('Relevancia autorías')
            d = row[1].add_paragraph(evaluator.authorship_str)

        if accreditation and accreditation.name:
            row = table.add_row().cells
            d = row[0].add_paragraph()
            d.add_run('Tipo acreditación')
            d = row[1].add_paragraph(accreditation.name)

        for row in table.rows:
            row.height = Cm(0.9)

    def print_scale(self, document, evaluation: SexenioEvaluation, bookmark:int):
        """
        Método que imprime la baremación obtenida tras evaluar la lista de artículos.
        :param document documento sobre el que se imprimirá la baremación
        :param evaluation objeto que contiene los datos de la evaluación del sexenio
        :param bookmark número asignado a la sección baremación en el documento
        """
        document.add_heading(str(bookmark) + '. Baremación mínima obtenida', level=1)
        p = document.add_paragraph('En la baremación que se ha realizado solo se ha tenido en cuenta los requisitos mínimos obligatorios ' +
                                    'de cada comité teniendo en cuenta los criterios de la ANECA.')

        if evaluation:
            if evaluation.punctuation > 0:
                p.add_run('\nLa puntuación mínima que podría obtener según los requisitos es: ' +
                          str(evaluation.punctuation) + '.')
            else:
                p.add_run(
                    '\nNo se ha podido obtener una baremación con las puntuaciones de la producción científica principal.')

            if evaluation.observation:
                p.add_run('\n'+evaluation.observation)

    def print_article(self, document, article, roposition, evaluator: Evaluator) -> bool:
        """
        Método encargado de consultar la información de un artículo e insertarlo en el informe
        :param document documento sobre el que se imprime el artículo
        :param article objeto que contiene la información del artículo
        :param roposition posición del artículo en el documento
        :param evaluator evaluador que ha realizado la evaluación de dicho artículo
        :return True si se ha incluido correctamente el artículo o False si no ha sido posible.
        """
        try:
            if article:
                document.add_heading('Posición '+str(roposition), level=2)
                npages = str(article.get_pages_number())
                document.add_heading(
                    'Información general del artículo', level=3)
                table = document.add_table(rows=0, cols=1)
                table.style = 'Table Grid'
                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('Revista: ').bold = True
                d.add_run(article.magazine+'\n')
                d.add_run('Editorial: ').bold = True
                d.add_run(article.editorial)
                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('Título: ').bold = True
                d.add_run(article.title+'\n')
                d.add_run('Autores: ').bold = True
                d.add_run(str(article.authors_number))

                if(evaluator.authors_alert != -1 and article.authors_number > evaluator.authors_alert):
                    d.add_run('\n')
                    eval = 'comité'
                    if evaluator.is_commission:
                        eval = 'comisión'

                    d.add_run('\nNúmero de autores ('+str(article.authors_number)+') por encima de lo recomendado para este ' +
                              eval + ' ('+str(evaluator.authors_alert) + ') , en este ' + eval +
                              ' se tiene muy en cuenta el número de autores, justificar muy bien la aportación y carga de trabajo en el artículo.').italic = True

                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('Año: ').bold = True
                d.add_run(article.get_year()+'\n')
                d.add_run('Volumen: ').bold = True
                d.add_run(article.volume+'\n')
                d.add_run('Número: ').bold = True
                d.add_run(str(article.number)+'\n')
                d.add_run('Número de páginas: ').bold = True
                d.add_run(npages+'\n')
                d.add_run('Página inicio: ').bold = True
                d.add_run(article.start_page+'\n')
                d.add_run('Página fin: ').bold = True
                d.add_run(article.end_page)

                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('ISSN: ').bold = True
                d.add_run(article.issn)

                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('DOI: ').bold = True
                d.add_run(article.doi)

                document.add_heading(
                    'Información JCR (Journal Citation Reports) del artículo', level=3)
                msg = 'Tiene un factor de impacto de ' + \
                    str(article.impact) + ' en JCR, año ' + \
                        article.get_year()+'.\n'
                msg += 'Ocupa la posición ' + str(article.position) + ' de un total de ' + \
                    str(article.num_magazines) + ' revistas. Por lo que está en el cuartil ' + \
                    article.quartile + 'º.'
                document.add_paragraph(msg)

                document.add_heading(
                    'Web of Science Colección Principal (WOS CC)', level=3)

                if article.wos_cites:
                    document.add_paragraph('Total Nº de citas: ' +
                                           article.wos_cites + '.')
                else:
                    document.add_paragraph('No se encontraron citas en WOS.')

                document.add_heading('SCOPUS', level=3)
                if article.scopus_cites:
                    document.add_paragraph('Total Nº de citas: ' +
                                           article.scopus_cites + '.')
                else:
                    document.add_paragraph(
                        'No se encontraron citas en SCOPUS.')

                document.add_heading('Semantic Scholar', level=3)
                if article.ss_cites:
                    document.add_paragraph('Total Nº de citas: ' +
                                           article.ss_cites + '.')
                else:
                    document.add_paragraph(
                        'No se encontraron citas en Semantic Scholar.')

            return True
        except Exception as e:
            print(e)
            self.notify_update(
                'ERROR en la inserción de un artículo en el informe.')
        return False

    def print_scientific_production(self, document, articles: list, evaluator, 
        production_type: str, bookmark:int):
        """
        Método que imprime la produccion cientifica de un investigador
        :param document documento sobre el que se imprimirá
        :param articles lista de elementos que forman la producción científica
        :param evaluator comisión que evalúa la acreditación
        :param production_type define el tipo de producción científica (principal, secundaria)
        :param bookmark número que determina la sección en el documento 
        """
        if production_type == 'principal':
            document.add_heading('2. Ranking de Producción Científica\n' + str(bookmark) + '. Producción principal', level=1)
        else:
            document.add_heading(str(bookmark) + '. Producción sustitutoria', level=1)
            document.add_paragraph(
                    'A continuación se detallan hasta 30 artículos sustitutorios por si el investigador desea sustituir producción científica calificada como principal:')
        
        if articles:
            roposition = 1
            for art in articles:
                if self.print_article(document, art, roposition, evaluator):
                    roposition += 1
        else:
            document.add_paragraph(
                'No se encontró producción científica ' + production_type + '.')
            self.notify_update(
                'No se encontró producción científica ' + production_type + '.')

    def print_books(self, document, books, evaluator, bookmark:int):
        """
        Método encargado de insertar los libros en el informe.
        :param document documento sobre el que se imprimirá
        :param books lista de elementos a imprimir
        :param evaluator comisión que evalúa la acreditación
        :param bookmark número que determina la sección en el documento 
        """
        document.add_heading(
            str(bookmark) + '. Libros y monografías científicas', level=1)

        if not books.empty:
            eval = 'el comité seleccionado'
            if evaluator.is_commission:
                eval = 'la comisión seleccionada'
            document.add_paragraph(
                'A continuación se detallan libros o monografías científicas que el autor ha llevado a cabo, ya que en ' + eval + ': '+evaluator.nombre+' pueden ser valorados.')
            for i in range(len(books)):
                document.add_paragraph(
                    '-Título: '+books.iloc[i]['titulo.value']+'. '+'Fecha de publicación: '+books.iloc[i]['anioFecha.value'], style='List Bullet')
        else:
            document.add_paragraph('No se encontraron libros.')
            self.notify_update(
                'No se encontraron libros')

    def print_chapter_books(self, document, chapters, evaluator, bookmark:int):
        """
        Método que inserta los capítulos de libros en el informe.
        :param document documento sobre el que se imprimirá
        :param chapters lista de elementos a imprimir
        :param evaluator comisión que evalúa la acreditación
        :param bookmark número que determina la sección en el documento 
        """
        document.add_heading(str(bookmark) + '. Capítulos de libros', level=1)

        if not chapters.empty:
            eval = 'el comité seleccionado'
            if evaluator.is_commission:
                eval = 'la comisión seleccionada'
            document.add_paragraph(
                'A continuación se detallan los capítulos de libros que el autor ha llevado a cabo, ya que en ' + eval + ': '+evaluator.name+' pueden ser valorados.')

            for i in range(len(chapters)):
                document.add_paragraph('-Título: '+chapters.iloc[i]['titulo.value']+'. ' +
                                       'Fecha de publicación: '+chapters.iloc[i]['anioFecha.value'], style='List Bullet')

        else:
            document.add_paragraph('No se encontraron capítulos de libros.')
            self.notify_update(
                'No se encontratron capítulos de libros.')

    def print_conferences(self, document, elements, bookmark:int):
        """
        Método encargado de insertar los trabajos presentados en congresos 
        nacionales e internacionales en el informe.
        :param document documento sobre el que se imprimirá
        :param elements lista de elementos a imprimir
        :param bookmark número que determina la sección en el documento 
        """
        document.add_heading(
            str(bookmark) + '. Trabajos presentados en congresos nacionales e internacionales', level=1)

        if elements:
            document.add_paragraph('En este informe se insertarán como máximo 20 trabajos presentados en congresos.')
            posicion = 1
            for element in elements:
                document.add_heading('Posición '+str(posicion), level=2)
                document.add_heading(
                    'Información general del trabajo:', level=3)
                table = document.add_table(rows=0, cols=1)
                table.style = 'Table Grid'
                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('Título: ').bold = True
                d.add_run(element.title+'\n')
                if element.authors:
                    d.add_run('Autores: ').bold = True
                    d.add_run(element.authors+'\n')
                    d.add_run('Número de autores: ').bold = True
                    d.add_run(str(element.num_authors))
                if element.date:
                    d.add_run('\nFecha publicación: ').bold = True
                    d.add_run(element.date)
                posicion += 1

        else:
            document.add_paragraph(
                'No se encontraron trabajos presentados en congresos.')
            self.notify_update(
                'No se encontraron trabajos presentados en congresos.')

    def print_patents(self, document, patents, bookmark:int):
        """
        Método encargado de insertar las patentes en el informe.
        :param document documento sobre el que se imprimirá
        :param patents lista de elementos a imprimir
        :param bookmark número que determina la sección en el documento 
        """
        document.add_heading(
            str(bookmark) + '. Patentes', level=1)

        if patents:
            document.add_paragraph('En este informe se insertarán como máximo 20 patentes.')

            for patente in patents:
                msg ='-Título: ' + patente.title
                if patente.date:                   
                    msg += ', Fecha de publicación: ' + patente.date
                if patente.authors:
                    msg += ', Autores: ' + patente.authors 
                if msg:
                    document.add_paragraph(msg + '.', style='List Bullet')
        else:
            document.add_paragraph('No se encontraron patentes.')
            self.notify_update(
                'No se encontraron patentes')

    def print_conferences_patents(self, document, evaluator:Evaluator, 
    conferences:list, patents:list, bookmark):
        if evaluator.conferences:
            self.notify_update('Insertando en el informe los trabajos presentados en congresos.')
            if conferences and len(conferences)>20:
                conferences = conferences[0:20]
            self.print_conferences(document, conferences, bookmark)

        if evaluator.patents:
            self.notify_update('Insertando en el informe las patentes.')
            if evaluator.conferences:
                bookmark +=1
            if len(patents)>20:
                patents = patents[0:20]
            self.print_patents(document, patents, bookmark)   

    def save_report(self, document, researcher_name, is_sexenio: bool) -> str:
        """
        Método que calcula el nombre del informe y lo almacena en la ubicación correspondiente. 
        :param document documento sobre el que se imprimirá
        :param researcher_name nombre del investigador
        :param is_sexenio define si es o no un sexenio el informe
        :return str nombre del documento guardado.
        """
        date_time = datetime.now().strftime("%m-%d-%Y-%H%M%S")
        doc_name = 'Informe'
        if is_sexenio:
            doc_name += 'Sexenios'
        else:
            doc_name += 'Acreditacion'

        doc_name += researcher_name+date_time+'.docx'
        doc_name = doc_name.replace(' ', '')
        document.save(doc_name)
        return doc_name

    def upload_file(self, filename, url_cdn):
        """
        Método encargado de subir el informe al CDN.
        :param filename nombre del archivo
        :para url_cdn url de CDN
        :return respuesta de la petición de subida al CDN.
        """
        response = None     
        if url_cdn:   
            with open('./'+filename, 'rb') as f:                
                try:
                    files = [('file', (filename, f, 'application/docx'))]
                    response = self.rpa.post(url_cdn, data_body={}, files=files)
                except Exception as e:
                    print(str(e))
                    self.notify_update('Error en la subida del informe al CDN.')            
                f.close()
        return response
