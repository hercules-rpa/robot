import os
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement, ns

class Report():
    def __init__(self, path=None):
        if not path:
            self.document = Document()
            parentdir = os.path.abspath(os.path.join(os.path.dirname(os.path.realpath(__file__)),os.pardir))
            self.document.add_picture(parentdir +"/Configurations/cabecera.png",width=Inches(5.9),height=Inches(1))
        else:
            self.document = Document(path)
        
        self.add_page_number(self.document.sections[0].footer.paragraphs[0].add_run("\t"))
        self.document.styles["Footer"]

    def create_element(self, name):
        """
        Creación de un elemento OXML.

        :param name str: nombre del elemento
        :return Elemento OXML creado.
        """
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        """
        Creación del atributo necesario para el documento.

        :param element
        :param name
        :param value
        """
        element.set(ns.qn(name), value)

    def add_page_number(self, run):
        """
        Método para añadir el número de página al documento.

        :param run: Instancia del documento.
        """
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def get_document(self):
        """
        Método que devuelve el documento creado.

        :return documento si existe.
        """
        if self.document:
            return self.document
        return None

    def add_title(self, titulo):
        """
        Método que añade el título que se le pasa por parámetro

        :param titulo str: Título del documento.
        """
        try:
            self.document.add_heading(titulo, 0)
        except Exception as e:
            print("Problema al insertar la cabecera: " + str(e))

    def save_document(self, path):
        self.document.save(path)

    def add_page_break(self):
        self.document.add_page_break()
